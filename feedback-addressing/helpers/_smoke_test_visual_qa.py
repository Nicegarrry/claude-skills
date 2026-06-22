#!/usr/bin/env python3
"""Smoke test for helpers/visual_qa.py.

Builds a 2-paragraph synthetic .docx, runs visual_qa.py, asserts a PDF + ≥1
PNG land on disk. If soffice is missing, the test asserts the helper returns
the structured `skipped` result and exits 0 (graceful).
"""
from __future__ import annotations

import json
import sys
import tempfile
from pathlib import Path

HERE = Path(__file__).resolve().parent
sys.path.insert(0, str(HERE))

import visual_qa  # noqa: E402

from docx import Document  # noqa: E402


def build_synthetic_docx(out_path: Path) -> None:
    doc = Document()
    doc.add_heading("Smoke test", level=1)
    doc.add_paragraph(
        "This is the first paragraph of a tiny synthetic document used by the "
        "visual-QA smoke test. It exists only to make sure soffice + pdftoppm "
        "produce a PDF and at least one PNG."
    )
    doc.add_paragraph(
        "Second paragraph. Lorem ipsum dolor sit amet. We force a small body "
        "so the PDF render is a single page."
    )
    doc.save(str(out_path))


def main() -> int:
    with tempfile.TemporaryDirectory() as tdir:
        d = Path(tdir)
        doc_path = d / "smoke.docx"
        out_dir = d / "vq"
        build_synthetic_docx(doc_path)
        result = visual_qa.run(doc_path, out_dir, dpi=100)
        print(json.dumps(result, indent=2))
        if result.get("skipped"):
            # graceful skip is acceptable
            print("[smoke] skipped (graceful):", result.get("reason"))
            return 0
        pdf = Path(result["pdf"])
        assert pdf.exists() and pdf.stat().st_size > 0, "PDF missing or empty"
        assert result["page_count"] >= 1, f"expected ≥1 PNG, got {result['page_count']}"
        for p in result["pages"]:
            assert Path(p).exists() and Path(p).stat().st_size > 0, f"PNG missing: {p}"
        print(f"[smoke] OK: pdf={pdf.name}, pages={result['page_count']}")
        return 0


if __name__ == "__main__":
    raise SystemExit(main())
