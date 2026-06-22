#!/usr/bin/env python3
"""
render_tables.py — render the reviewer-facing simplified table from the
internal `02-internal-table.md` (SKILL.md step 11).

WHAT IT DOES
------------
- Parses the 11-column internal markdown table (`feedback_id`, `source`,
  `feedback_summary`, `tier`, `plan_to_address`, `knowledge_gap?`,
  `research_dispatched?`, `clarifying_Q?`, `change_applied`,
  `change_location`, `status`).
- Renders one of three output shapes:
    * email-draft   — full markdown email block (salutation + cover + 3-col
                      table + sign-off). Default.
    * markdown-table — just the 3-col simplified table.
    * docx-export   — write a .docx of the simplified table.

WHAT IT DOESN'T DO
------------------
- Does NOT mutate the internal table.
- Does NOT auto-send anything; output is text/file only.

USAGE
-----
    render_tables.py <internal-table.md>
        --format {email-draft|markdown-table|docx-export}
        --doc-stem STEM
        [--reviewer-name NAME]
        [--signoff NAME]      # name to sign the reviewer email with (your name)
        [--out PATH]
"""
from __future__ import annotations

import argparse
import re
import sys
from pathlib import Path
from typing import Any


# ----- markdown table parser ------------------------------------------------
def parse_md_table(md_text: str) -> tuple[list[str], list[dict[str, str]]]:
    """
    Parse the FIRST markdown pipe-table found in `md_text`.
    Returns (header_cols, rows-as-dicts).
    """
    lines = md_text.splitlines()
    table_lines: list[str] = []
    in_table = False
    for ln in lines:
        if ln.lstrip().startswith("|"):
            table_lines.append(ln.strip())
            in_table = True
        elif in_table and not ln.strip():
            break
        elif in_table and not ln.lstrip().startswith("|"):
            break
    if len(table_lines) < 2:
        raise SystemExit("no markdown table found in input")

    def cells(row: str) -> list[str]:
        # split on | but keep escaped \|
        parts = row.strip().strip("|").split("|")
        return [p.strip() for p in parts]

    header = cells(table_lines[0])
    # second line is the separator (---|---|...)
    body_rows: list[dict[str, str]] = []
    for ln in table_lines[2:]:
        if re.match(r"^\|?\s*[-:]+\s*(\|\s*[-:]+\s*)*\|?$", ln):
            continue
        c = cells(ln)
        if len(c) != len(header):
            # pad / truncate
            if len(c) < len(header):
                c = c + [""] * (len(header) - len(c))
            else:
                c = c[: len(header)]
        body_rows.append(dict(zip(header, c)))
    return header, body_rows


# ----- 3-col simplified rows ------------------------------------------------
def to_simplified(rows: list[dict[str, str]]) -> list[dict[str, str]]:
    out: list[dict[str, str]] = []
    for r in rows:
        feedback = r.get("feedback_summary", "").strip()
        status = r.get("status", "").strip().lower()
        change = r.get("change_applied", "").strip()
        where = r.get("change_location", "").strip()

        if status == "applied":
            what = change if change and change != "—" else "Addressed."
        elif status == "escalated":
            what = "No change yet, confirming with author."
        elif status == "skipped-no-change-needed":
            what = "No change, see note."
            where = where if where and where != "—" else "—"
        else:
            what = change if change and change != "—" else "—"

        out.append(
            {
                "Feedback": feedback,
                "What changed": what,
                "Where": where if where else "—",
            }
        )
    return out


# ----- renderers ------------------------------------------------------------
def render_markdown_table(simplified: list[dict[str, str]]) -> str:
    if not simplified:
        return "_(no items)_\n"
    cols = ["Feedback", "What changed", "Where"]
    lines = ["| " + " | ".join(cols) + " |", "|" + "|".join(["---"] * len(cols)) + "|"]
    for row in simplified:
        cells = [(row.get(c, "") or "").replace("|", "\\|").replace("\n", " ") for c in cols]
        lines.append("| " + " | ".join(cells) + " |")
    return "\n".join(lines) + "\n"


def render_email_draft(
    simplified: list[dict[str, str]],
    doc_stem: str,
    reviewer_name: str | None,
    signoff: str | None = None,
) -> str:
    salutation = f"Hi {reviewer_name}," if reviewer_name else "Hi,"
    table_md = render_markdown_table(simplified)
    n_total = len(simplified)
    n_applied = sum(1 for r in simplified if r["What changed"] not in ("No change yet, confirming with author.", "No change, see note.", "—"))
    n_escalated = sum(1 for r in simplified if r["What changed"] == "No change yet, confirming with author.")

    cover = (
        f"Thanks for the feedback on **{doc_stem}**. I've worked through your comments and "
        f"the table below summarises what changed and where. {n_applied} of {n_total} items "
        "are addressed in this revision; "
    )
    if n_escalated:
        cover += f"{n_escalated} I'd like to confirm with you before applying. "
    cover += "Tracked changes are visible in the attached document."

    closing = "Happy to discuss any of these, just reply on the doc or here.\n\n"
    signoff_block = f"Thanks,\n{signoff}\n" if signoff else "Thanks,\n"
    return (
        f"{salutation}\n\n"
        f"{cover}\n\n"
        f"{table_md}\n"
        f"{closing}"
        f"{signoff_block}"
    )


def render_docx_export(simplified: list[dict[str, str]], out_path: Path) -> bool:
    try:
        from docx import Document  # type: ignore
    except ImportError:
        print("[render_tables] docx-export requires python-docx — skipping", file=sys.stderr)
        return False
    doc = Document()
    doc.add_heading("Responses to feedback", level=1)
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Feedback"
    hdr[1].text = "What changed"
    hdr[2].text = "Where"
    for row in simplified:
        cells = table.add_row().cells
        cells[0].text = row.get("Feedback", "")
        cells[1].text = row.get("What changed", "")
        cells[2].text = row.get("Where", "")
    doc.save(out_path)
    return True


# ----- main -----------------------------------------------------------------
def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("internal_table", type=Path)
    ap.add_argument("--format", choices=["email-draft", "markdown-table", "docx-export"], default="email-draft")
    ap.add_argument("--doc-stem", default="this document")
    ap.add_argument("--reviewer-name", default=None)
    ap.add_argument("--signoff", default=None, help="name to sign the reviewer email with")
    ap.add_argument("--out", type=Path, default=None)
    args = ap.parse_args()

    md = args.internal_table.read_text(encoding="utf-8")
    _, rows = parse_md_table(md)
    simplified = to_simplified(rows)

    if args.format == "markdown-table":
        out_text = render_markdown_table(simplified)
        if args.out:
            args.out.parent.mkdir(parents=True, exist_ok=True)
            args.out.write_text(out_text, encoding="utf-8")
            print(f"[render_tables] wrote markdown table → {args.out}", file=sys.stderr)
        else:
            print(out_text)
    elif args.format == "email-draft":
        out_text = render_email_draft(simplified, args.doc_stem, args.reviewer_name, args.signoff)
        if args.out:
            args.out.parent.mkdir(parents=True, exist_ok=True)
            args.out.write_text(out_text, encoding="utf-8")
            print(f"[render_tables] wrote email draft → {args.out}", file=sys.stderr)
        else:
            print(out_text)
    elif args.format == "docx-export":
        if args.out is None:
            print("--out PATH required for docx-export", file=sys.stderr)
            return 2
        ok = render_docx_export(simplified, args.out)
        if not ok:
            return 3
        print(f"[render_tables] wrote .docx → {args.out}", file=sys.stderr)

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
